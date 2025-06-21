# src/parsers/doc_parser.py
from pathlib import Path
from typing import List
from .base_parser import BaseParser
from src.utils.exceptions import InvalidFileError, FileAccessError
from docx import Document
import subprocess
import os
import tempfile
import shutil


class DOCParser(BaseParser):
    """
    Парсер документов Microsoft Word (.doc, .docx)
    Поддерживаемые форматы: .docx, .doc
    """

    SUPPORTED_MIME_TYPES = ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']

    def __init__(self, file_path: str):
        self._temp_file_created = None  # 用于记录临时生成的 docx 路径
        if file_path.lower().endswith(".doc"):
            file_path = self._convert_doc_to_docx(file_path)
        super().__init__(file_path)

    def _convert_doc_to_docx(self, file_path: str) -> str:
        with tempfile.TemporaryDirectory() as tmpdirname:
            try:
                subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'docx', file_path,
                    '--outdir', tmpdirname
                ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

                base_name = Path(file_path).stem + ".docx"
                converted_path = Path(tmpdirname) / base_name

                if not converted_path.exists():
                    raise InvalidFileError("doc", Path(file_path), "ошибка преобразования в .docx")

                final_path = Path(file_path).with_suffix(".docx")
                shutil.copyfile(converted_path, final_path)

                self._temp_file_created = final_path
                return str(final_path)

            except Exception as e:
                raise InvalidFileError("doc", Path(file_path), "не удалось конвертировать .doc в .docx") from e

    def _validate_file(self):
        super()._validate_file()
        try:
            Document(self.file_path)
        except Exception as e:
            raise InvalidFileError("docx", Path(self.file_path), "файл поврежден") from e

    def extract_text(self) -> str:
        try:
            doc = Document(self.file_path)
            return '\n'.join(para.text for para in doc.paragraphs).strip()
        except Exception as e:
            raise FileAccessError(Path(self.file_path), "чтение текста") from e

    def extract_metadata(self) -> dict:
        return {}

    def extract_images(self) -> List[bytes]:
        return []

    def __del__(self):
        if self._temp_file_created and os.path.exists(self._temp_file_created):
            try:
                os.remove(self._temp_file_created)
            except Exception:
                pass
