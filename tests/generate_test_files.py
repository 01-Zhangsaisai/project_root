import os
import random
import string
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

def random_string(length=10):
    """Генерация случайной строки заданной длины"""
    letters = string.ascii_letters + string.digits
    return ''.join(random.choice(letters) for i in range(length))

def create_valid_pdf(file_path):
    c = canvas.Canvas(file_path, pagesize=letter)
    c.drawString(100, 750, f"Пример корректного PDF-файла: {random_string()}.")
    c.save()

def create_invalid_pdf(file_path):
    with open(file_path, 'wb') as f:
        f.write(b'\x00\x00\x00')  # Запись поврежденных данных

def create_valid_doc(file_path):
    with open(file_path, 'w') as f:
        f.write(f"Пример корректного DOC-файла: {random_string()}.\n")

def create_invalid_doc(file_path):
    with open(file_path, 'wb') as f:
        f.write(b'\x00\x00\x00')

def create_valid_docx(file_path):
    doc = Document()
    doc.add_heading(f'Пример DOCX-файла: {random_string()}', level=1)
    doc.add_paragraph(f'Это корректный пример DOCX-файла: {random_string()}.')
    doc.save(file_path)

def create_invalid_docx(file_path):
    with open(file_path, 'wb') as f:
        f.write(b'\x00\x00\x00')

def create_valid_html(file_path):
    with open(file_path, 'w') as f:
        f.write(f"<html><body><h1>Пример корректного HTML-файла: {random_string()}.</h1></body></html>")

def create_invalid_html(file_path):
    with open(file_path, 'w') as f:
        f.write("<html><body><h1>Некорректный HTML-файл")


def create_invalid_djvu(file_path):
    with open(file_path, 'wb') as f:
        f.write(b'\x00\x00\x00')

def generate_test_files(num_files=5):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    valid_dir = os.path.join(base_dir, "data", "valid")
    invalid_dir = os.path.join(base_dir, "data", "invalid")

    os.makedirs(valid_dir, exist_ok=True)
    os.makedirs(invalid_dir, exist_ok=True)

    for i in range(num_files):
        create_valid_pdf(os.path.join(valid_dir, f"sample_{i + 1}_{random_string(5)}.pdf"))
        create_valid_doc(os.path.join(valid_dir, f"sample_{i + 1}_{random_string(5)}.doc"))
        create_valid_docx(os.path.join(valid_dir, f"sample_{i + 1}_{random_string(5)}.docx"))
        create_valid_html(os.path.join(valid_dir, f"sample_{i + 1}_{random_string(5)}.html"))

        create_invalid_pdf(os.path.join(invalid_dir, f"corrupted_{i + 1}_{random_string(5)}.pdf"))
        create_invalid_doc(os.path.join(invalid_dir, f"corrupted_{i + 1}_{random_string(5)}.doc"))
        create_invalid_docx(os.path.join(invalid_dir, f"corrupted_{i + 1}_{random_string(5)}.docx"))
        create_invalid_html(os.path.join(invalid_dir, f"corrupted_{i + 1}_{random_string(5)}.html"))
        create_invalid_djvu(os.path.join(invalid_dir, f"corrupted_{i + 1}_{random_string(5)}.djvu"))

if __name__ == "__main__":
    generate_test_files(num_files=10)