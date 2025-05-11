# src/parsers/docx_parser.py
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pathlib import Path
from .base_parser import BaseParser
from src.utils.exceptions import InvalidFileError


class DOCXParser(BaseParser):
    """
    Парсер документов Microsoft Word (.docx)
    Поддерживаемые форматы: application/vnd.openxmlformats-officedocument.wordprocessingml.document, .docx
    """

    SUPPORTED_MIME_TYPES = ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']
    SUPPORTED_EXTENSIONS = ['.docx']

    def _validate_file(self):
        """Специфическая проверка DOCX"""
        super()._validate_file()
        path = Path(self.file_path)
        try:
            Document(path)
        except PackageNotFoundError:
            raise InvalidFileError("docx", path, "файл поврежден") from None

    def extract_text(self) -> str:
        """Извлечение текста"""
        doc = Document(self.file_path)
        return "\n".join(para.text for para in doc.paragraphs)

    def extract_metadata(self) -> dict:
        """Извлечение свойств документа"""
        doc = Document(self.file_path)
        return {
            "title": doc.core_properties.title,
            "author": doc.core_properties.author,
            "created": doc.core_properties.created
        }

    def extract_images(self) -> list:
        """Извлечение встроенных изображений"""
        doc = Document(self.file_path)
        return [rel.target_part.blob for rel in doc.part.rels.values() if "image" in rel.reltype]